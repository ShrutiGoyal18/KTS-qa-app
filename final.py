import os
import sys
import uuid
import re
import torch
import hashlib
import openai
import requests
import numpy as np
import faiss
import pickle
from tqdm import tqdm
from typing import List, Dict, Tuple
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
from transformers import AutoTokenizer, AutoModelForSequenceClassification

from langchain.memory import ConversationBufferMemory
from langchain.schema import messages_from_dict, messages_to_dict
from langchain_openai import AzureChatOpenAI

# --------------------------
# Load environment vars
# --------------------------
# Replace the current dotenv loading code at line ~15-30
import os
from dotenv import load_dotenv

# Try to load from Streamlit secrets first, then fall back to .env file
try:
    import streamlit as st
    # Use Streamlit secrets if available
    AZURE_OPENAI_API_KEY = st.secrets.get("AZURE_OPENAI_API_KEY", None)
    AZURE_ENDPOINT = st.secrets.get("AZURE_OPENAI_ENDPOINT", None)
    AZURE_DEPLOYMENT_NAME = st.secrets.get("AZURE_DEPLOYMENT_NAME", None)
    
    if AZURE_OPENAI_API_KEY and AZURE_ENDPOINT and AZURE_DEPLOYMENT_NAME:
        print("Using Azure OpenAI credentials from Streamlit secrets")
    else:
        # Fall back to .env file
        load_dotenv()
        AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
        AZURE_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT") 
        AZURE_DEPLOYMENT_NAME = os.getenv("AZURE_DEPLOYMENT_NAME")
except ImportError:
    # Not running in Streamlit environment
    load_dotenv()
    AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
    AZURE_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT") 
    AZURE_DEPLOYMENT_NAME = os.getenv("AZURE_DEPLOYMENT_NAME")

# Print values for debugging (remove in production)
print(f"Azure Endpoint: {AZURE_ENDPOINT}")
print(f"Azure Deployment: {AZURE_DEPLOYMENT_NAME}")
print(f"API Key: {'Set' if AZURE_OPENAI_API_KEY else 'Not set'}")

# Check if key variables are missing
if not AZURE_OPENAI_API_KEY or not AZURE_ENDPOINT or not AZURE_DEPLOYMENT_NAME:
    print("Warning: Some Azure OpenAI environment variables are missing.")
    print("Please check your .env file.")

# ---------------------
# Memory Setup
# ---------------------
memory = ConversationBufferMemory(return_messages=True)

# ---------------------
# Load and Chunk Documents
# ---------------------

def doc_hash(path: str) -> str:
    """Generate a hash for the document to use as ID"""
    with open(path, 'rb') as f:
        return hashlib.md5(f.read()).hexdigest()

def extract_text_from_docx(doc_path):
    """Extract text from Word document"""
    try:
        import docx
        doc = docx.Document(doc_path)
        
        # Extract text from paragraphs
        paragraphs_text = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(" | ".join(row_text))
        
        # Combine all text
        full_text = "\n".join(paragraphs_text + tables_text)
        print(f"Extracted {len(paragraphs_text)} paragraphs and {len(tables_text)} table rows")
        return full_text
        
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"Error extracting text from Word document: {str(e)}")
        sys.exit(1)

def load_doc_chunks(doc_path: str, chunk_size=1000,overlap=150)->List[str]:
    """Load and chunk text from a document"""
    
    # Check file extension
    file_ext = os.path.splitext(doc_path)[1].lower()
    
    if file_ext == '.docx':
        # Extract from DOCX
        full_text = extract_text_from_docx(doc_path)
    else:
        print(f"Unsupported file format: {file_ext}")
        print("Please use a .docx file")
        return []
    
    # Clean and chunk the text
    if full_text:
        print(f"Successfully extracted {len(full_text)} characters from the document")
        full_text = re.sub(r'\s+', ' ', full_text)
        words = full_text.split()
        print(f"Found {len(words)} words to chunk")
        
        chunks = []
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            chunks.append(chunk)
        
        print(f"Created {len(chunks)} text chunks")
        return chunks
    else:
        print("No text was extracted from the document.")
        return []

# -------------------------
# FAISS Index Management
# -------------------------

class FAISSIndex:
    def __init__(self, dimension=384):
        self.dimension = dimension
        self.index = faiss.IndexFlatL2(dimension)  # Using L2 distance
        self.metadata = {}
        self.id_counter = 0
        self.id_to_idx = {}  # Maps external IDs to index positions
    
    def save(self, path):
        """Save the FAISS index and metadata to disk"""
        os.makedirs(os.path.dirname(path), exist_ok=True)  # Ensure directory exists
        faiss.write_index(self.index, f"{path}.index")
        with open(f"{path}.metadata", 'wb') as f:
            pickle.dump({
                'metadata': self.metadata,
                'id_counter': self.id_counter,
                'id_to_idx': self.id_to_idx
            }, f)
    
    def load(self, path):
        """Load the FAISS index and metadata from disk"""
        if os.path.exists(f"{path}.index") and os.path.exists(f"{path}.metadata"):
            self.index = faiss.read_index(f"{path}.index")
            with open(f"{path}.metadata", 'rb') as f:
                data = pickle.load(f)
                self.metadata = data['metadata']
                self.id_counter = data['id_counter']
                self.id_to_idx = data['id_to_idx']
            return True
        return False
    
    def upsert(self, vectors_with_metadata):
        """Add vectors to the index with metadata"""
        vectors = []
        ids = []
        
        for id_or_none, vector, metadata in vectors_with_metadata:
            if id_or_none is None:
                id_str = str(self.id_counter)
                self.id_counter += 1
            else:
                id_str = id_or_none
            
            vectors.append(np.array(vector, dtype=np.float32))
            self.metadata[id_str] = metadata
            ids.append(id_str)
        
        if vectors:
            vectors_array = np.array(vectors, dtype=np.float32)
            faiss.normalize_L2(vectors_array)  # Normalize for cosine similarity
            
            start_idx = self.index.ntotal
            self.index.add(vectors_array)
            
            # Update ID mapping
            for i, id_str in enumerate(ids):
                self.id_to_idx[id_str] = start_idx + i
    
    def query(self, vector, top_k=10, filter=None):
        """Search for similar vectors"""
        vector = np.array([vector], dtype=np.float32)
        faiss.normalize_L2(vector)  # Normalize for cosine similarity
        
        distances, indices = self.index.search(vector, top_k)
        
        matches = []
        for i, idx in enumerate(indices[0]):
            if idx != -1:  # Valid index
                # Find ID for this index position
                id_str = None
                for id_val, pos in self.id_to_idx.items():
                    if pos == idx:
                        id_str = id_val
                        break
                
                if id_str and id_str in self.metadata:
                    metadata = self.metadata[id_str]
                    
                    # Apply filter if specified
                    if filter:
                        if "doc_id" in filter and "doc_id" in metadata:
                            if filter["doc_id"]["$eq"] != metadata["doc_id"]:
                                continue
                    
                    matches.append({
                        "id": id_str,
                        "score": float(1.0 - distances[0][i]),  # Convert to similarity score
                        "metadata": metadata
                    })
        
        return {"matches": matches}

# -------------------------
# Embeddings & Indexing
# -------------------------

def embed_and_store(chunks: List[str], doc_id: str, model_name='BAAI/bge-large-en-v1.5'):
    if not chunks:
        print("No chunks to embed.")
        return None
    
    # Create paths
    index_path = os.path.join("faiss_indexes", doc_id)
    os.makedirs("faiss_indexes", exist_ok=True)
    
    # Check if index files exist early before loading model
    if os.path.exists(f"{index_path}.index") and os.path.exists(f"{index_path}.metadata"):
        print("Found existing index files for this document.")
        
        # Load index directly
        index = FAISSIndex()
        if index.load(index_path):
            # Simple verification - just check if we have any vectors with this doc_id
            count = 0
            for metadata in index.metadata.values():
                if "doc_id" in metadata and metadata["doc_id"] == doc_id:
                    count += 1
                    if count > 0:  # We found at least one chunk
                        print(f"Document already indexed with {count} chunks. Skipping embedding.")
                        return index
    
    # If we reach here, we need to create/update the index
    print("Creating new embeddings for document...")
    model = SentenceTransformer(model_name)
    dimension = model.get_sentence_embedding_dimension()
    
    # Create or load FAISS index
    index = FAISSIndex(dimension=dimension)
    index.load(index_path)  # Try loading, but it's ok if it doesn't exist
    
    batch = []
    for chunk in tqdm(chunks, desc="Embedding chunks"):
        # For BGE models, adding a query prefix improves retrieval performance
        vector = model.encode(chunk).tolist()
        metadata = {"text": chunk, "doc_id": doc_id}
        batch.append((str(uuid.uuid4()), vector, metadata))
    
    index.upsert(batch)
    index.save(index_path)
    print(f"{len(batch)} chunks embedded and stored in FAISS index.")
    return index

# -------------------------
# Reranker (BGE)
# -------------------------

def rerank(query: str, passages: List[str], top_k=5, model_name='BAAI/bge-reranker-base'):
    """Rerank passages using a simplified approach to avoid compatibility issues"""
    if not passages:
        return []
        
    # Ensure top_k isn't larger than the number of passages
    top_k = min(top_k, len(passages))
    
    if top_k == 0:
        return []
    
    try:
        # Use sentence-transformers cross-encoder instead
        from sentence_transformers import CrossEncoder
        
        # Load cross-encoder model
        model = CrossEncoder(model_name)
        
        # Create score pairs
        sentence_pairs = [[query, passage] for passage in passages]
        
        # Get scores
        scores = model.predict(sentence_pairs)
        
        # Get indices of top scoring passages
        top_indices = np.argsort(scores)[-top_k:][::-1]  # Sort and get top-k indices
        
        # Return reranked passages
        reranked_passages = [passages[idx] for idx in top_indices]
        return reranked_passages
        
    except Exception as e:
        print(f"Error in reranking: {str(e)}")
        # If reranking fails, return original passages
        return passages[:top_k]

# -------------------------
# Query with Memory
# -------------------------

# Make sure your Azure deployment is correctly specified
def ask_gpt_azure(query: str, context: List[str]) -> str:
    if not context:
        return "I couldn't find any relevant information to answer your question."
        
    prompt = f"""You are a highly intelligent assistant. Use the provided context to answer the user's question as accurately and completely as possible.
###
Carefully follow these instructions:
###
1. Analyze the entire document context and identify all relevant pieces of information related to the question.
2. If relevant information is found, summarize and synthesize it to form a detailed and accurate response.
3. If the context contains structured data (like tables or lists), extract and format the data clearly.
4. If the context has no relevant content, you **must** include this sentence in your answer:  
   **"The answer you are looking for is not found in the uploaded document!"**
5. Consider the current question and any context from the previous chat history to provide a coherent and context-aware response.

--- Context ---
{chr(10).join(context)}
--- End Context ---

Question: {query}
Answer:"""

    memory.save_context({"input": query}, {"output": "..."})  # stub to simulate memory update
    
    try:
        # Create LangChain Azure OpenAI chat model
        chat = AzureChatOpenAI(
            azure_endpoint=AZURE_ENDPOINT, 
            api_key=AZURE_OPENAI_API_KEY,
            api_version="2024-02-01",
            temperature=0,
            max_tokens=500,
            azure_deployment=AZURE_DEPLOYMENT_NAME
        )
        
        # Create messages array with system, memory, and user content
        messages = [{"role": "system", "content": "You are a knowledgeable assistant."}]
        messages.extend(memory.chat_memory.messages)
        messages.append({"role": "user", "content": prompt})
        
        # Get response from LangChain model
        response = chat.invoke(messages)
        
        # Extract content from response
        answer = response.content
        
        memory.save_context({"input": query}, {"output": answer})
        return answer
    
    except Exception as e:
        print(f"Error calling Azure OpenAI: {str(e)}")
        
        # Fallback to using local response with context
        context_text = "\n".join(context[:3])  # Show just the first 3 chunks to avoid overwhelming output
        return f"""
I encountered an issue connecting to the AI service. Here's the most relevant context I found:

{context_text}

Please check your Azure OpenAI configuration. The error was: {str(e)}
"""

# -------------------------
# Chatbot Interface
# -------------------------

def chat_with_bot(query, index, embed_model_name='BAAI/bge-large-en-v1.5', reranker_top_k=5):
    if index is None:
        return "No document has been properly indexed. Please check the document loading process."
        
    embed_model = SentenceTransformer(embed_model_name)
    # For BGE models, adding a query prefix improves retrieval performance
    query_text = "Represent this sentence for searching relevant passages: " + query
    query_vec = embed_model.encode(query_text).tolist()

    # Query the index with more results to give the reranker more options
    res = index.query(vector=query_vec, top_k=10)  # Increased to 10 from 5
    
    # Check if we have any results
    if not res['matches']:
        return "I don't have any information about that. The document might not have been properly indexed."
    
    retrieved_chunks = [match['metadata']['text'] for match in res['matches']]

    # Apply reranking
    print("Reranking retrieved chunks...")
    try:
        reranked_chunks = rerank(query, retrieved_chunks, top_k=reranker_top_k)
        print(f"Reranked {len(reranked_chunks)} chunks")
        
        # Use the reranked chunks
        answer = ask_gpt_azure(query, reranked_chunks)
    except Exception as e:
        print(f"Reranking failed: {str(e)}, falling back to vector search results")
        # Fallback to original chunks if reranking fails
        answer = ask_gpt_azure(query, retrieved_chunks[:reranker_top_k])
    
    return answer

# -------------------------
# Main Execution
# -------------------------

if __name__ == "__main__":
    # Check if document path is provided as command line argument
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
    else:
        # Default location - update with your DOCX file
        doc_path = r"C:\Users\U6077339\OneDrive - Clarivate Analytics\KTS_vs_python\KTS Docketing - US Patent V 3.8.2.docx" 
    
    # Strip quotes if they were included in the path
    doc_path = doc_path.strip('"\'')
    
    if not os.path.exists(doc_path):
        print(f"Error: Document file not found at {doc_path}")
        print("Please provide the correct path to the DOCX file.")
        sys.exit(1)
    
    # Check if it's a DOCX file
    file_ext = os.path.splitext(doc_path)[1].lower()
    if file_ext != '.docx':
        print(f"Error: File {doc_path} is not a Word document (.docx)")
        print("Please provide a valid .docx file.")
        sys.exit(1)
    
    doc_id = doc_hash(doc_path)

    print("Loading and chunking document...")
    chunks = load_doc_chunks(doc_path)
    
    if not chunks:
        print("Error: Failed to extract any text chunks from the document.")
        sys.exit(1)
    
    print("Creating FAISS index and storing embeddings (if not already done)...")
    faiss_index = embed_and_store(chunks, doc_id=doc_id)
    
    print("🤖 Ready! Ask your question:")
    while True:
        user_query = input("You: ")
        if user_query.lower() in ['exit', 'quit', 'q']:
            break
        response = chat_with_bot(user_query, faiss_index)
        print(f"\nBot: {response}\n")